"""
IP Prefix Extractor Module

This module extracts IP prefixes from various file formats including:
- Text files (.txt, .log, .conf, .ini)
- CSV files (.csv)
- Excel files (.xls, .xlsx, .xlsm)
- Word documents (.docx)
- PDF files (.pdf)
- JSON files (.json)
- XML files (.xml, .html, .htm)

"""

import re
import os
import csv
import traceback
import argparse
import pandas as pd
import chardet
from openpyxl import load_workbook
try:
    from docx import Document
except ImportError:
    # If python-docx is not installed, make the function return empty results
    def extract_from_docx(file_path):
        print(f"Error: python-docx module not installed. Please install it with 'pip install python-docx'")
        return [], ['IP_Address']
try:
    import pdfplumber
except ImportError:
    print("Warning: pdfplumber module not installed. PDF extraction will be limited.")
    print("Please install it with 'pip install pdfplumber'")
    pdfplumber = None
import json
import xml.etree.ElementTree as ET

# Add PyPDF2 for enhanced PDF extraction
try:
    import PyPDF2
except ImportError:
    print("Warning: PyPDF2 module not installed. PDF extraction will be limited.")
    print("Please install it with 'pip install PyPDF2'")
    PyPDF2 = None

def detect_encoding(file_path):
    """Detect the encoding of a file."""
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read())
    return result['encoding']

def extract_ips_from_text(text):
    """Extract IP prefixes from text using regex."""
    # Regular expression for IPv4 prefixes (like 192.168.1.0/24)
    ipv4_prefix_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}/\d{1,2}\b'
    
    # Regular expression for IPv6 prefixes
    ipv6_prefix_pattern = r'\b(?:[0-9a-fA-F]{1,4}:){1,7}[0-9a-fA-F]{1,4}/\d{1,3}\b'
    
    # Add debug info
    print(f"Processing text of length: {len(text)}")
    print(f"First 100 characters: {text[:100]}")
    
    # Extract IP prefixes
    ipv4_prefixes = re.findall(ipv4_prefix_pattern, text)
    ipv6_prefixes = re.findall(ipv6_prefix_pattern, text)
    
    # Print results for debugging
    print(f"Found {len(ipv4_prefixes)} IPv4 prefixes and {len(ipv6_prefixes)} IPv6 prefixes")
    if ipv4_prefixes:
        print(f"Sample IPv4 prefixes: {ipv4_prefixes[:5]}")
    
    # Check if text might contain IP prefixes with alternate formatting
    if len(ipv4_prefixes) <= 1:
        # Try different variations for IP prefix formats
        alt_ipv4_pattern = r'(?:\d{1,3}[._]){3}\d{1,3}[/_-]\d{1,2}'
        alt_ipv4 = re.findall(alt_ipv4_pattern, text)
        if alt_ipv4:
            print(f"Found {len(alt_ipv4)} potential alternate format IP prefixes")
            # Convert to standard format
            standard_format = [re.sub(r'[._]', '.', ip).replace('_', '/').replace('-', '/') for ip in alt_ipv4]
            ipv4_prefixes.extend(standard_format)
    
    # Deduplicate
    return list(set(ipv4_prefixes + ipv6_prefixes))

def extract_from_txt(file_path):
    """Extract IP addresses from text files."""
    try:
        encoding = detect_encoding(file_path)
        print(f"Detected encoding for {file_path}: {encoding}")
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as file:
            content = file.read()
        
        print(f"File size: {len(content)} characters")
        
        lines = content.split('\n')
        print(f"Total lines in file: {len(lines)}")
        
        # Print a sample of the file
        print("File sample (first 3 lines):")
        for i in range(min(3, len(lines))):
            print(f"Line {i+1}: {lines[i][:100]}{'...' if len(lines[i]) > 100 else ''}")
        
        # Try to identify column headers from the first line
        headers = []
        if lines and not re.search(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', lines[0]):
            potential_headers = lines[0].split()
            if potential_headers:
                headers = potential_headers
                print(f"Potential headers found: {headers}")
        
        # Extract all IP addresses from the content
        ips = extract_ips_from_text(content)
        print(f"Total unique IP addresses found: {len(ips)}")
        
        # Process the file line by line to extract IP addresses with context
        line_by_line_results = []
        
        for i, line in enumerate(lines):
            if not line.strip():
                continue
                
            ips_in_line = extract_ips_from_text(line)
            if ips_in_line:
                # For each IP in the line, create a separate entry
                for ip in ips_in_line:
                    # Get some context around the IP (the whole line)
                    entry = {'IP_Address': ip, 'Line_Number': i+1, 'Context': line.strip()}
                    line_by_line_results.append(entry)
        
        print(f"Line-by-line extraction found {len(line_by_line_results)} IP entries")
        
        # If we found results line by line, use those
        if line_by_line_results:
            # Determine headers - always include IP_Address and Line_Number
            result_headers = ['IP_Address', 'Line_Number', 'Context']
            if headers:
                # Try to extract values for the headers from each line
                for entry in line_by_line_results:
                    line_parts = entry['Context'].split()
                    for i, header in enumerate(headers):
                        if i < len(line_parts):
                            entry[header] = line_parts[i]
                            if header not in result_headers:
                                result_headers.append(header)
            
            return line_by_line_results, result_headers
        
        # If no line-by-line results, fall back to simple list of IPs
        if not line_by_line_results and ips:
            result = [{'IP_Address': ip} for ip in ips]
            return result, ['IP_Address']
        
        # If no IPs were found at all
        return [], ['IP_Address']
    except Exception as e:
        print(f"Error processing text file {file_path}: {e}")
        return [], ['IP_Address']

def extract_from_csv(file_path):
    """Extract IP addresses from CSV files."""
    try:
        # Try to detect encoding
        encoding = detect_encoding(file_path)
        
        # Try different CSV dialect options
        for delimiter in [',', ';', '\t', '|']:
            try:
                df = pd.read_csv(file_path, delimiter=delimiter, encoding=encoding, error_bad_lines=False, warn_bad_lines=True, quoting=csv.QUOTE_MINIMAL)
                break
            except Exception as e:
                continue
        else:
            # If all delimiters fail, try the pandas auto-detection
            df = pd.read_csv(file_path, encoding=encoding, sep=None, engine='python', error_bad_lines=False)
        
        # Get headers
        headers = list(df.columns)
        
        # Find columns containing IP addresses
        ip_containing_rows = []
        
        # First, log some information about the DataFrame
        print(f"DataFrame shape: {df.shape}")
        print(f"DataFrame headers: {headers}")
        
        # Convert DataFrame to string to search for IPs
        df_str = df.astype(str)
        
        # Search each row for IP addresses
        for idx, row in df_str.iterrows():
            row_dict = {}
            ip_found = False
            row_ips = []
            
            # First, search for IPs across the whole row text
            row_text = ' '.join(row.values)
            all_ips_in_row = extract_ips_from_text(row_text)
            
            # Then process each cell
            for col in headers:
                cell_value = row[col]
                ips_in_cell = extract_ips_from_text(cell_value)
                
                if ips_in_cell:
                    ip_found = True
                    for ip in ips_in_cell:
                        row_dict[f"{col}_{len(row_ips) + 1}" if len(row_ips) > 0 else col] = ip
                        row_ips.append(ip)
                else:
                    # Only include non-IP columns if we found at least one IP
                    if ip_found or all_ips_in_row:
                        row_dict[col] = row[col]
            
            # If no IPs found in individual cells but in the whole row, add them
            if not ip_found and all_ips_in_row:
                for i, ip in enumerate(all_ips_in_row):
                    col_name = f"IP_Address_{i+1}" if i > 0 else "IP_Address"
                    row_dict[col_name] = ip
                ip_found = True
            
            if ip_found:
                ip_containing_rows.append(row_dict)
                
        print(f"Found {len(ip_containing_rows)} rows containing IP addresses")
        
        if not ip_containing_rows:
            # If no IPs found in individual cells, try searching all text
            all_text = ' '.join([' '.join(df_str[col]) for col in df_str.columns])
            ips = extract_ips_from_text(all_text)
            ip_containing_rows = [{'IP_Address': ip} for ip in ips]
            headers = ['IP_Address']
            
        return ip_containing_rows, headers
    except Exception as e:
        print(f"Error processing CSV file {file_path}: {e}")
        # Try treating it as a plain text file
        return extract_from_txt(file_path)

def extract_from_excel(file_path):
    """Extract IP addresses from Excel files."""
    try:
        # Use pandas to read the Excel file
        df = pd.read_excel(file_path, sheet_name=None)
        
        all_ips = []
        headers = []
        
        # Process each sheet
        for sheet_name, sheet_df in df.items():
            sheet_headers = list(sheet_df.columns)
            
            if not headers:
                headers = sheet_headers
            
            # Convert all cells to string for IP searching
            sheet_df = sheet_df.astype(str)
            
            # Search each row for IP addresses
            for _, row in sheet_df.iterrows():
                row_dict = {}
                ip_found = False
                
                for col in sheet_headers:
                    cell_value = row[col]
                    ips_in_cell = extract_ips_from_text(str(cell_value))
                    
                    if ips_in_cell:
                        ip_found = True
                        row_dict[col] = ips_in_cell[0]  # Take the first IP in the cell
                    else:
                        row_dict[col] = cell_value
                
                if ip_found:
                    all_ips.append(row_dict)
        
        if not all_ips:
            # Try a different approach - extract all text and find IPs
            workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
            all_text = ""
            
            for sheet in workbook:
                for row in sheet.iter_rows():
                    row_text = ' '.join([str(cell.value) if cell.value is not None else '' for cell in row])
                    all_text += row_text + ' '
            
            ips = extract_ips_from_text(all_text)
            all_ips = [{'IP_Address': ip} for ip in ips]
            headers = ['IP_Address']
            
        return all_ips, headers
    except Exception as e:
        print(f"Error processing Excel file {file_path}: {e}")
        return [], ['IP_Address']

def extract_from_docx(file_path):
    """Extract IP addresses from Word documents."""
    try:
        doc = Document(file_path)
        all_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Check if the document has tables
        has_tables = len(doc.tables) > 0
        
        # If document has tables, try to extract headers and structured data
        if has_tables:
            headers = []
            structured_data = []
            
            for table in doc.tables:
                if not headers and len(table.rows) > 0:
                    # Extract headers from the first row
                    headers = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
                
                # Process each row (skip header row if headers were found)
                start_row = 1 if headers else 0
                for i in range(start_row, len(table.rows)):
                    row = table.rows[i]
                    row_data = [cell.text.strip() for cell in row.cells]
                    
                    # Check if this row contains any IP addresses
                    row_text = ' '.join(row_data)
                    ips_in_row = extract_ips_from_text(row_text)
                    
                    if ips_in_row:
                        if headers:
                            row_dict = {headers[j]: value for j, value in enumerate(row_data) if j < len(headers)}
                            structured_data.append(row_dict)
                        else:
                            structured_data.append({'IP_Address': ips_in_row[0]})
            
            if structured_data:
                return structured_data, headers or ['IP_Address']
        
        # If no structured data from tables, or no tables, extract IPs from text
        ips = extract_ips_from_text(all_text)
        
        if ips:
            # Look for potential headers near the IPs
            lines = all_text.split('\n')
            potential_headers = []
            
            for i, line in enumerate(lines):
                if any(ip in line for ip in ips) and i > 0:
                    prev_line = lines[i-1].strip()
                    if prev_line and not any(extract_ips_from_text(prev_line)):
                        potential_headers = re.split(r'\s{2,}', prev_line)
                        break
            
            if potential_headers:
                # Try to map IPs to the headers
                result = []
                for ip in ips:
                    # Find the line containing this IP
                    for line in lines:
                        if ip in line:
                            # Split the line by whitespace
                            parts = re.split(r'\s{2,}', line)
                            
                            if len(parts) == len(potential_headers):
                                entry = {potential_headers[i]: parts[i] for i in range(len(parts))}
                                result.append(entry)
                            else:
                                result.append({potential_headers[0] if potential_headers else 'IP_Address': ip})
                            break
                    else:
                        result.append({potential_headers[0] if potential_headers else 'IP_Address': ip})
                
                return result, potential_headers
            else:
                return [{'IP_Address': ip} for ip in ips], ['IP_Address']
        else:
            return [], ['IP_Address']
    except Exception as e:
        print(f"Error processing Word file {file_path}: {e}")
        return [], ['IP_Address']

def extract_from_pdf(file_path):
    """Extract IP prefixes from PDF files with enhanced methods for high volume extraction."""
    # Add PyPDF2 extraction method
    def extract_with_pypdf2():
        """Alternative PDF extraction using PyPDF2"""
        if PyPDF2 is None:
            return set()
            
        print("Method 5: Extracting with PyPDF2...")
        extracted_prefixes = set()
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                print(f"PyPDF2 found {page_count} pages")
                
                for i in range(page_count):
                    page = reader.pages[i]
                    text = page.extract_text()
                    
                    if text:
                        prefixes = extract_ips_from_text(text)
                        if prefixes:
                            print(f"  PyPDF2 page {i+1}: Found {len(prefixes)} IP prefixes")
                            extracted_prefixes.update(prefixes)
        except Exception as e:
            print(f"PyPDF2 extraction error (non-fatal): {e}")
            
        return extracted_prefixes
    try:
        print(f"\nEnhanced PDF extraction for {file_path}")
        all_text = ""
        all_ips = set()  # Use a set to avoid duplicates
        result = []
        
        # Method 1: Using pdfplumber for text extraction
        try:
            print("Method 1: Extracting with pdfplumber...")
            with pdfplumber.open(file_path) as pdf:
                print(f"PDF has {len(pdf.pages)} pages")
                
                # Process each page
                for page_num, page in enumerate(pdf.pages, 1):
                    print(f"Processing page {page_num}...")
                    
                    # Extract text
                    page_text = page.extract_text() or ""
                    all_text += page_text + "\n"
                    
                    # Find IPs in the page text
                    page_ips = extract_ips_from_text(page_text)
                    if page_ips:
                        print(f"  Found {len(page_ips)} IPs in page {page_num} text")
                        all_ips.update(page_ips)
                    
                    # Try various extraction methods for tables
                    tables = page.extract_tables()
                    if tables:
                        print(f"  Found {len(tables)} tables in page {page_num}")
                        for table_idx, table in enumerate(tables):
                            if not table:
                                continue
                                
                            table_text = ' '.join([' '.join([str(cell) if cell else '' for cell in row]) for row in table])
                            table_ips = extract_ips_from_text(table_text)
                            
                            if table_ips:
                                print(f"    Table {table_idx+1}: Found {len(table_ips)} IPs")
                                all_ips.update(table_ips)
                                
                                # Try to extract structured data from the table
                                headers = [str(cell).strip() for cell in table[0] if cell]
                                if headers:
                                    for row_idx, row in enumerate(table[1:], 1):
                                        row_dict = {}
                                        row_has_ip = False
                                        
                                        for col_idx, cell in enumerate(row):
                                            if col_idx < len(headers) and cell:
                                                cell_str = str(cell)
                                                cell_ips = extract_ips_from_text(cell_str)
                                                
                                                if cell_ips:
                                                    for ip_idx, ip in enumerate(cell_ips):
                                                        header = headers[col_idx]
                                                        ip_header = f"{header}_{ip_idx+1}" if ip_idx > 0 else header
                                                        row_dict[ip_header] = ip
                                                        row_has_ip = True
                                                else:
                                                    row_dict[headers[col_idx]] = cell_str
                                        
                                        if row_has_ip:
                                            # Add metadata
                                            row_dict['Page'] = page_num
                                            row_dict['Table'] = table_idx + 1
                                            row_dict['Row'] = row_idx
                                            result.append(row_dict)
        except Exception as e:
            print(f"pdfplumber method error (non-fatal): {e}")
        
        print(f"After pdfplumber extraction: Found {len(all_ips)} unique IPs")
        
        # Method 2: Process text line by line to catch more IPs
        lines = all_text.split('\n')
        print(f"Processing {len(lines)} lines from PDF text...")
        for line_num, line in enumerate(lines, 1):
            if not line.strip():
                continue
                
            line_ips = extract_ips_from_text(line)
            if line_ips:
                for ip in line_ips:
                    entry = {
                        'IP_Address': ip,
                        'Line': line_num,
                        'Context': line.strip()
                    }
                    result.append(entry)
                    all_ips.add(ip)
        
        print(f"After line-by-line processing: Found {len(all_ips)} unique IPs")
        
        # Method 3: Special case for lists of IPs
        if len(all_ips) < 10:  # If we didn't find many IPs, try another approach
            print("Trying special extraction for lists of IPs...")
            # Look for patterns like sequences of numbers that might be IPs
            ip_candidates = re.findall(r'\b\d{1,3}[\s\.,_]+\d{1,3}[\s\.,_]+\d{1,3}[\s\.,_]+\d{1,3}\b', all_text)
            
            if ip_candidates:
                print(f"Found {len(ip_candidates)} potential IP candidates")
                for candidate in ip_candidates:
                    # Try to normalize to standard IP format
                    normalized = re.sub(r'[\s,_]', '.', candidate)
                    if re.match(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', normalized):
                        all_ips.add(normalized)
                        result.append({'IP_Address': normalized, 'Source': 'Special Extraction'})
        
        print(f"After special extraction: Found {len(all_ips)} unique IPs")
        
        # Method 4: Raw character analysis for very problematic PDFs
        if len(all_ips) < 10:
            print("Trying raw character analysis...")
            # Try to find sequences that look like IPs but might have unusual separators
            raw_text = all_text.replace('\n', ' ')
            potential_ips = []
            
            i = 0
            while i < len(raw_text) - 6:  # Need at least 7 chars for an IP
                # Check if we have a digit
                if raw_text[i].isdigit():
                    # Look ahead for a pattern like: digits.digits.digits.digits
                    j = i + 1
                    dots = 0
                    segments = 1
                    segment_valid = True
                    current_segment = raw_text[i]
                    
                    while j < len(raw_text) and dots < 4 and segments <= 4:
                        if raw_text[j].isdigit() and segment_valid:
                            current_segment += raw_text[j]
                            # Check if segment is valid (0-255)
                            if len(current_segment) > 3 or int(current_segment) > 255:
                                segment_valid = False
                        elif (raw_text[j] in '., \t_-') and segment_valid and current_segment:
                            if segments < 4:  # Only count separators between segments
                                dots += 1
                                segments += 1
                                current_segment = ""
                                segment_valid = True
                            else:
                                break  # We've reached the end of a potential IP
                        else:
                            break  # Invalid character for an IP
                        
                        j += 1
                    
                    # Check if we found a valid pattern
                    if dots == 3 and segments == 4 and segment_valid:
                        ip_candidate = raw_text[i:j].strip()
                        # Normalize the IP format
                        normalized = re.sub(r'[, \t_-]', '.', ip_candidate)
                        if re.match(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', normalized):
                            potential_ips.append(normalized)
                i += 1
            
            # Add unique IPs from this method
            unique_potential_ips = set(potential_ips)
            for ip in unique_potential_ips:
                if ip not in all_ips:
                    all_ips.add(ip)
                    result.append({'IP_Address': ip, 'Source': 'Raw Analysis'})
        
        # Method 5: Use PyPDF2 as well
        pypdf2_ips = extract_with_pypdf2()
        if pypdf2_ips:
            print(f"PyPDF2 found {len(pypdf2_ips)} IPs, {len(pypdf2_ips - all_ips)} are new")
            for ip in pypdf2_ips - all_ips:
                all_ips.add(ip)
                result.append({'IP_Address': ip, 'Source': 'PyPDF2'})
        
        print(f"Final count: Found {len(all_ips)} unique IPs across all methods")
        
        # Determine final result format
        if result:
            # Get all possible headers
            all_headers = set()
            for entry in result:
                all_headers.update(entry.keys())
            
            # Ensure IP_Address is the first header
            headers = ['IP_Address']
            for header in sorted(list(all_headers)):
                if header != 'IP_Address':
                    headers.append(header)
                    
            print(f"Using headers: {headers}")
            return result, headers
        elif all_ips:
            # If we have IPs but no structured data, create simple entries
            simple_result = [{'IP_Address': ip} for ip in all_ips]
            return simple_result, ['IP_Address']
        else:
            print("No IP addresses found in PDF")
            return [], ['IP_Address']
            
    except Exception as e:
        print(f"Overall PDF extraction error: {str(e)}")
        traceback_info = traceback.format_exc()
        print(f"Traceback: {traceback_info}")
        return [], ['IP_Address']

def extract_from_json(file_path):
    """Extract IP addresses from JSON files."""
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding) as file:
            data = json.load(file)
        
        # Convert data to string to search for IPs
        data_str = json.dumps(data)
        ips = extract_ips_from_text(data_str)
        
        # Try to find header information
        headers = []
        structured_data = []
        
        def process_json_item(item, path=""):
            nonlocal headers, structured_data
            
            if isinstance(item, dict):
                ip_found = False
                row_dict = {}
                
                for key, value in item.items():
                    current_path = f"{path}.{key}" if path else key
                    
                    if isinstance(value, (dict, list)):
                        process_json_item(value, current_path)
                    else:
                        value_str = str(value)
                        ips_in_value = extract_ips_from_text(value_str)
                        
                        if ips_in_value:
                            ip_found = True
                            row_dict[key] = ips_in_value[0]
                            if key not in headers:
                                headers.append(key)
                        elif ip_found:  # Only add other fields if we found an IP
                            row_dict[key] = value_str
                            if key not in headers:
                                headers.append(key)
                
                if ip_found and row_dict:
                    structured_data.append(row_dict)
                    
            elif isinstance(item, list):
                for i, element in enumerate(item):
                    current_path = f"{path}[{i}]"
                    process_json_item(element, current_path)
        
        process_json_item(data)
        
        if structured_data:
            return structured_data, headers
        else:
            return [{'IP_Address': ip} for ip in ips], ['IP_Address']
    except Exception as e:
        print(f"Error processing JSON file {file_path}: {e}")
        # Try treating it as a plain text file
        return extract_from_txt(file_path)

def extract_from_xml(file_path):
    """Extract IP addresses from XML files."""
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding) as file:
            content = file.read()
        
        # First, just extract all IPs from the text
        ips = extract_ips_from_text(content)
        
        # Try to parse as XML
        try:
            root = ET.fromstring(content)
            
            # Search for elements containing IP addresses
            structured_data = []
            headers = set()
            
            def process_element(element, path=""):
                nonlocal structured_data, headers
                
                # Check if this element has text with an IP
                if element.text and element.text.strip():
                    ips_in_text = extract_ips_from_text(element.text)
                    if ips_in_text:
                        # Use the element's tag or path as header
                        header = element.tag
                        headers.add(header)
                        
                        # Create a dictionary for this item
                        item_dict = {header: ips_in_text[0]}
                        
                        # Add attributes and child elements as additional fields
                        for attr_name, attr_value in element.attrib.items():
                            item_dict[f"{header}_{attr_name}"] = attr_value
                            headers.add(f"{header}_{attr_name}")
                        
                        for child in element:
                            if child.text and child.text.strip():
                                item_dict[child.tag] = child.text.strip()
                                headers.add(child.tag)
                        
                        structured_data.append(item_dict)
                
                # Process child elements
                for child in element:
                    process_element(child, f"{path}/{element.tag}")
            
            process_element(root)
            
            if structured_data:
                return structured_data, list(headers)
        except ET.ParseError:
            pass  # Not valid XML, continue with text extraction
        
        # Fall back to plain text extraction
        return [{'IP_Address': ip} for ip in ips], ['IP_Address']
    except Exception as e:
        print(f"Error processing XML file {file_path}: {e}")
        # Try treating it as a plain text file
        return extract_from_txt(file_path)

def save_to_csv(data, headers, output_file):
    """Save extracted data to a CSV file with only IP_Address field."""
    try:
        with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
            # Only use the IP_Address field
            writer = csv.writer(csvfile)
            writer.writerow(['IP_Address'])  # Single column header
            
            # Write only the IP addresses
            for row in data:
                if isinstance(row, dict) and 'IP_Address' in row:
                    writer.writerow([row['IP_Address']])
                elif isinstance(row, list) and len(row) > 0:
                    writer.writerow([row[0]])
                else:
                    writer.writerow([row])
        
        print(f"Results saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error saving to CSV: {e}")
        return False

def extract_ips(file_path, output_file=None):
    """
    Extract IP prefixes from a file based on its extension.
    Returns simplified data with only IP_Address.
    
    Args:
        file_path (str): Path to the file to extract IPs from
        output_file (str, optional): Path to save the CSV results
        
    Returns:
        tuple: (extracted_data, headers, success_flag)
    """
    # This portion remains the same as the original - extract IPs using various methods
    if not os.path.exists(file_path):
        print(f"Error: The file {file_path} was not found.")
        return [], [], False
    
    print(f"\n{'='*50}")
    print(f"Processing file: {file_path}")
    print(f"{'='*50}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    print(f"File extension: {file_extension}")
    
    # First, try a quick check for IP prefixes using text mode
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            sample = file.read(10000)  # Read first 10KB for a quick check
            quick_ip_check = extract_ips_from_text(sample)
            print(f"Quick IP prefix check in first 10KB: {len(quick_ip_check)} prefixes found")
    except Exception as e:
        # This error is common for binary files, so no need to print it
        pass
    
    # Determine the appropriate extraction method based on file extension
    print(f"Using extraction method based on file extension: {file_extension}")
    
    if file_extension in ['.txt', '.log', '.conf', '.ini']:
        data, headers = extract_from_txt(file_path)
    elif file_extension == '.csv':
        data, headers = extract_from_csv(file_path)
    elif file_extension in ['.xls', '.xlsx', '.xlsm']:
        data, headers = extract_from_excel(file_path)
    elif file_extension == '.docx':
        data, headers = extract_from_docx(file_path)
    elif file_extension == '.pdf':
        data, headers = extract_from_pdf(file_path)
    elif file_extension == '.json':
        data, headers = extract_from_json(file_path)
    elif file_extension in ['.xml', '.html', '.htm']:
        data, headers = extract_from_xml(file_path)
    else:
        # Try as text file for unsupported extensions
        print(f"Unsupported file extension: {file_extension}. Trying as text file.")
        data, headers = extract_from_txt(file_path)
    
    # If the specialized method didn't find IP prefixes, try as plain text
    if not data:
        print("No IP prefixes found with specialized method. Trying as plain text...")
        data, headers = extract_from_txt(file_path)
    
    # Last resort: Try binary mode and look for IP prefix patterns
    if not data:
        print("Still no IP prefixes found. Trying binary mode...")
        try:
            with open(file_path, 'rb') as file:
                binary_content = file.read()
                text_content = binary_content.decode('utf-8', errors='replace')
                prefixes = extract_ips_from_text(text_content)
                if prefixes:
                    data = [{'IP_Address': prefix} for prefix in prefixes]
                    headers = ['IP_Address']
        except Exception as e:
            print(f"Binary mode error (non-fatal): {e}")
    
    # Simplify the data to only include IP_Address
    simplified_data = []
    for item in data:
        if isinstance(item, dict) and 'IP_Address' in item:
            simplified_data.append({'IP_Address': item['IP_Address']})
    
    # If we couldn't extract structured data but found raw IPs, use those
    if not simplified_data and data:
        if isinstance(data[0], dict):
            # Try to find any IP address field
            for item in data:
                for key, value in item.items():
                    if re.search(r'\b(?:\d{1,3}\.){3}\d{1,3}/\d{1,2}\b', str(value)):
                        simplified_data.append({'IP_Address': value})
                        break
        else:
            # Assume data is a list of IP addresses
            simplified_data = [{'IP_Address': ip} for ip in data if isinstance(ip, str)]
    
    # Generate output file name if not provided
    if not output_file and simplified_data:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_file = f"{base_name}_extracted_prefixes.csv"
    
    # Save results if output_file is provided
    success = False
    if simplified_data:
        print(f"Final result: Found {len(simplified_data)} entries containing IP prefixes")
        simplified_headers = ['IP_Address']
        if output_file:
            success = save_to_csv(simplified_data, simplified_headers, output_file)
        else:
            success = True
        return simplified_data, simplified_headers, success
    else:
        print("No IP prefixes found in the file.")
        return [], ['IP_Address'], False

def main():
    parser = argparse.ArgumentParser(description='Extract IP prefixes from various file formats.')
    parser.add_argument('file_path', help='Path to the input file')
    parser.add_argument('-o', '--output', help='Path to the output CSV file')
    
    args = parser.parse_args()
    
    data, headers, success = extract_ips(args.file_path, args.output)
    return success

if __name__ == "__main__":
    main()